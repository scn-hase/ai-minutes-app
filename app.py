from google.oauth2 import service_account
# app.py の先頭に追加
import io
from docx import Document
import vertexai
from vertexai.generative_models import GenerativeModel, Part
# app.py の先頭に追加
import os
import datetime
from google.cloud import storage
# app.py
import streamlit as st

# アプリのタイトルを設定
st.title("AI議事録作成アプリ 📄✍️")
st.markdown("""
このアプリは、音声ファイルをアップロードするだけで、AIが自動で文字起こしと議事録の作成を行います。
Gemini 2.5 Flashの強力な性能をぜひ体験してください！
""")

# ファイルアップロードのウィジェットを作成
uploaded_file = st.file_uploader(
    "議事録を作成したい音声ファイル（MP3, WAV）をアップロードしてください。",
    type=["mp3", "wav"]
)

    # ファイルがアップロードされたら処理を開始
if uploaded_file is not None:
    # --- この行以降は、すべて同じインデントレベルで始める ---
    
    st.success(f"ファイル「{uploaded_file.name}」がアップロードされました。")

    # --- 認証ブロック ---
    try:
        # Streamlit CloudのSecretsから認証情報を読み込む
        creds_dict = st.secrets["gcp_service_account"]
        creds = service_account.Credentials.from_service_account_info(creds_dict)
        storage_client = storage.Client(credentials=creds)
        project_id = "gizirokuapp"
        location = "asia-northeast1" # リージョンを東京に設定
        vertexai.init(project=project_id, location=location, credentials=creds)
    
    except (FileNotFoundError, KeyError):
        # ローカル環境の場合
        st.info("ローカル環境として実行します。")
        storage_client = storage.Client()
        project_id = "gizirokuapp"
        location = "asia-northeast1" # リージョンを東京に設定
        vertexai.init(project=project_id, location=location)
    
    # --- ファイル名の生成 ---
    # `with`ブロックの前に移動して、ロジックを明確にする
    timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    blob_name = f"{timestamp}-{uploaded_file.name}"
    
    # --- ファイルアップロードブロック ---
    # `try`ブロックと同じインデントレベルに修正する
    with st.spinner("ファイルをクラウドにアップロード中..."):
        
        bucket_name = "scn-giziroku-tokyo"
        bucket = storage_client.bucket(bucket_name)
        
        # GCSにファイルをアップロード
        blob = bucket.blob(blob_name)
        blob.upload_from_file(uploaded_file, content_type=uploaded_file.type)
        
        st.info(f"ファイルのアップロードが完了しました: gs://{bucket_name}/{blob_name}")
        
        # GCS上のファイルURIを後で使うために保存
        gcs_uri = f"gs://{bucket_name}/{blob_name}"

        # gcs_uri を取得した後の処理
if "gcs_uri" in locals():
    with st.spinner("AIが音声を文字起こし中です... この処理には数分かかることがあります。"):
       
        # 使用するモデルを指定 (Gemini 1.5 Flash)
        model = GenerativeModel(model_name="gemini-1.5-flash-001")
        
        # GCS上の音声ファイルを指定
        audio_file = Part.from_uri(
            mime_type=uploaded_file.type,
            uri=gcs_uri
        )
        
        # プロンプト（AIへの指示）を作成
        prompt = "この音声ファイルを日本語で文字起こししてください。"
        
        # AIにリクエストを送信
        response = model.generate_content([audio_file, prompt])
        
        # 結果を取得
        transcribed_text = response.text
        
        st.subheader("文字起こし結果")
        st.write(transcribed_text)
        
        # ここに議事録生成処理を続けて書く
    # transcribed_text を取得した後の処理
if "transcribed_text" in locals():
    with st.spinner("AIが議事録を生成中です..."):
        # 議事録生成用のプロンプトを作成
        # 指示を具体的に書くのがコツです（プロンプトエンジニアリング）
        prompt_for_minutes = f"""
        以下の会議の文字起こしテキストを元に、プロフェッショナルな議事録を作成してください。
        
        以下のフォーマットに従って、要点を明確にまとめてください。
        
        # 議事録
        
        ## 1. 会議の要約
        （会議全体のサマリーを3〜5行で記述）
        
        ## 2. 決定事項
        （会議で決定された事項を箇条書きでリストアップ）
        - 決定事項1
        - 決定事項2
        
        ## 3. ToDoリスト（担当者と期限）
        （発生したタスクを箇条書きでリストアップし、誰がいつまでに行うかを明記）
        - [ ] タスク1（担当：〇〇さん、期限：YYYY-MM-DD）
        - [ ] タスク2（担当：△△さん、期限：YYYY-MM-DD）
        
        ---
        
        # 文字起こしテキスト
        
        {transcribed_text}
        """

        # 再びAIにリクエストを送信
        response_minutes = model.generate_content(prompt_for_minutes)
        generated_minutes = response_minutes.text

        st.subheader("生成された議事録")
        st.markdown(generated_minutes)

        # ここにWordファイル出力処理を続けて書く    
        # generated_minutes を取得した後の処理
if "generated_minutes" in locals():
    with st.spinner("Wordファイルを生成中です..."):
        # Wordドキュメントを作成
        document = Document()
        
        document.add_heading('AI自動生成議事録', 0)
        
        # 生成された議事録を追加
        document.add_heading('生成された議事録', level=1)
        # Markdown形式のテキストを解析して、見出しやリストをWordに追加
        for line in generated_minutes.split('\n'):
            if line.startswith('### '):
                document.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('## '):
                document.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('# '):
                document.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('- '):
                document.add_paragraph(line, style='List Bullet')
            else:
                document.add_paragraph(line)
        
        document.add_page_break()
        
        # 元の文字起こしテキストを追加
        document.add_heading('文字起こし全文', level=1)
        document.add_paragraph(transcribed_text)
        
        # Wordファイルをメモリ上に保存
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        
    st.success("議事録の生成が完了しました！以下からダウンロードできます。")

    # ダウンロードボタンを表示
    st.download_button(
        label="議事録をWordファイルでダウンロード",
        data=file_stream,
        file_name=f"議事録_{uploaded_file.name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )